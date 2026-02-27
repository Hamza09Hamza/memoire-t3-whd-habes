from docx import Document
doc = Document('main_word.docx')

# Check all unique styles used
styles = {}
for p in doc.paragraphs:
    s = p.style.name
    if s not in styles:
        styles[s] = []
    if len(styles[s]) < 3 and p.text.strip():
        styles[s].append(p.text.strip()[:60])

print("Styles used in document:")
for s, examples in sorted(styles.items()):
    print(f"\n  [{s}] ({sum(1 for p in doc.paragraphs if p.style.name == s)} paragraphs)")
    for e in examples:
        print(f"    -> {e}")

# Check RTL in paragraph properties
print("\n\nRTL check on first few paragraphs:")
from docx.oxml.ns import qn
for p in doc.paragraphs[:5]:
    ppr = p._element.find(qn('w:pPr'))
    bidi = ppr.find(qn('w:bidi')) if ppr is not None else None
    text = p.text.strip()[:50] if p.text.strip() else "(empty)"
    print(f"  bidi={'yes' if bidi is not None else 'no'} | {text}")
