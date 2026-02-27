#!/usr/bin/env python3
"""
Build the final DOCX from scratch:
  1. Title page — matching the PDF reference layout, with our thesis data
  2. شكر وتقدير — our content from acknowledgments.tex, styled like the PDF
  3. الإهداء — our content from dedication.tex, styled like the PDF
  4. Body — from main_word_styled.docx, starting from المقدمة العامة
  5. Page numbering starts at 1 from the introduction
  6. Footnotes for references (per-page)
"""

import os, copy, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
BODY_DOCX = os.path.join(BASE, 'main_word_styled.docx')
OUTPUT = os.path.join(BASE, 'main_word_final.docx')

AF = 'Traditional Arabic'
LF = 'Times New Roman'

# ─── helpers ────────────────────────────────────────────────

def _rf(run, fa=AF, fl=LF, sz=14, bold=False, color=None, italic=False):
    """Set font on run."""
    rpr = run._element.get_or_add_rPr()
    f = OxmlElement('w:rFonts')
    f.set(qn('w:cs'), fa); f.set(qn('w:ascii'), fl); f.set(qn('w:hAnsi'), fl)
    for o in rpr.findall(qn('w:rFonts')): rpr.remove(o)
    rpr.insert(0, f)
    run.font.size = Pt(sz)
    sc = OxmlElement('w:szCs'); sc.set(qn('w:val'), str(int(sz*2)))
    for o in rpr.findall(qn('w:szCs')): rpr.remove(o)
    rpr.append(sc)
    run.font.bold = bold
    if bold:
        if rpr.find(qn('w:bCs')) is None: rpr.append(OxmlElement('w:bCs'))
    if italic:
        run.font.italic = italic
    if color: run.font.color.rgb = color
    if rpr.find(qn('w:rtl')) is None: rpr.append(OxmlElement('w:rtl'))

def _bidi(p):
    """Set bidi on paragraph."""
    ppr = p._element.get_or_add_pPr()
    if ppr.find(qn('w:bidi')) is None: ppr.append(OxmlElement('w:bidi'))

def _ap(doc, txt, sz=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
        sb=0, sa=0, color=None, ls=1.5, italic=False):
    """Add paragraph."""
    p = doc.add_paragraph()
    p.alignment = align; pf = p.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa); pf.line_spacing = ls
    _bidi(p)
    if txt:
        r = p.add_run(txt)
        _rf(r, sz=sz, bold=bold, color=color, italic=italic)
    return p

def _box(p, sz=16):
    """Box border around paragraph."""
    ppr = p._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    for s in ['top','bottom','left','right']:
        e = OxmlElement(f'w:{s}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz))
        e.set(qn('w:space'),'6'); e.set(qn('w:color'),'000000')
        bd.append(e)
    ppr.append(bd)

def _deco(p, char='✦', count=20, sz=12):
    """Add decorative line."""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bidi(p)
    r = p.add_run(char * count)
    _rf(r, sz=sz, color=RGBColor(0x99,0x99,0x99))

def _hr(p, sz=6, color='999999'):
    """Bottom border line on paragraph."""
    ppr = p._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),color)
    bd.append(b)
    ppr.append(bd)

def _page_border(sec):
    """Add page border to a section (like the PDF reference)."""
    sp = sec._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '24')
        el.set(qn('w:color'), '000000')
        pgBorders.append(el)
    sp.append(pgBorders)

# ─── front pages ────────────────────────────────────────────

def build_title_page(doc):
    _ap(doc, 'الجمهورية الجزائرية الديمقراطية الشعبية', sz=14, bold=True, sa=4)
    _ap(doc, 'وزارة التكوين والتعليم المهنيين', sz=14, bold=True, sa=8)

    # Decorative line
    d = doc.add_paragraph(); _deco(d, '═', 30, 14)

    _ap(doc, 'المعهد الوطني المتخصص في التكوين المهني للتسيير بالبليدة',
        sz=16, bold=True, sb=8, sa=24)

    _ap(doc, 'مذكــــــــــــرة نهايــــــــــــة التكويــــــــــــن',
        sz=20, bold=True, sb=24, sa=6)
    _ap(doc, 'للحصول على شهادة أهلية التقني السامــي', sz=14, bold=True, sa=4)
    _ap(doc, 'تخصص وكالة سياحة واسفار', sz=16, bold=True, sa=20)

    _ap(doc, 'بعنوان :', sz=14, bold=True, sb=12, sa=10)

    # Title in box
    tp = _ap(doc, 'تحديات وكالات الأسفار في ظل المنافسة مع المنصات الإلكترونية',
             sz=22, bold=True, sb=10, sa=10)
    _box(tp)

    # Host institution
    hp = _ap(doc, '', sz=14, bold=True, sb=20, sa=20)
    r1 = hp.add_run('مؤسسة الإستقبال : ')
    _rf(r1, sz=14, bold=True)
    r2 = hp.add_run('Le Grand Monarque LGM')
    _rf(r2, fa=LF, sz=14, bold=True)

    # Students / Supervisor table (borderless, RTL)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tp_ = tbl._element.find(qn('w:tblPr'))
    if tp_ is None:
        tp_ = OxmlElement('w:tblPr'); tbl._element.insert(0, tp_)
    # no borders
    tb = OxmlElement('w:tblBorders')
    for s in ['top','bottom','left','right','insideH','insideV']:
        e = OxmlElement(f'w:{s}')
        e.set(qn('w:val'),'none'); e.set(qn('w:sz'),'0')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'auto')
        tb.append(e)
    tp_.append(tb)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'),'5000'); tw.set(qn('w:type'),'pct')
    for o in tp_.findall(qn('w:tblW')): tp_.remove(o)
    tp_.append(tw)
    tp_.append(OxmlElement('w:bidiVisual'))

    # col 0 = right in RTL = students
    c = tbl.cell(0,0); c.paragraphs[0].clear()
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _bidi(p)
    r = p.add_run('إعداد المتربصين :'); _rf(r, sz=14, bold=True)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; _bidi(p2)
    r = p2.add_run('بن جيلالي صهيب'); _rf(r, sz=14)
    p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; _bidi(p3)
    r = p3.add_run('منذر ادريس مداني'); _rf(r, sz=14)

    # col 1 = left in RTL = supervisor
    c2 = tbl.cell(0,1); c2.paragraphs[0].clear()
    p4 = c2.paragraphs[0]; p4.alignment = WD_ALIGN_PARAGRAPH.CENTER; _bidi(p4)
    r = p4.add_run('إشراف الأستاذ :'); _rf(r, sz=14, bold=True)
    p5 = c2.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.CENTER; _bidi(p5)
    r = p5.add_run('لحرش أمير'); _rf(r, sz=14)

    # Decorative line
    _ap(doc, '', sz=6, sb=24, sa=0)
    d2 = doc.add_paragraph(); _deco(d2, '═', 30, 14)

    _ap(doc, 'دفعة : 2025 / 2026', sz=16, bold=True, sb=12, sa=0)


def build_acknowledgments(doc):
    """شكر وتقدير page — our content from acknowledgments.tex, styled like PDF."""
    # Page break
    pb = doc.add_paragraph()
    pb.paragraph_format.page_break_before = True
    _bidi(pb)

    # Decorative top
    d = doc.add_paragraph(); _deco(d, '❊', 15, 14)

    # Title
    _ap(doc, 'شكر وتقدير', sz=28, bold=True, sb=24, sa=8,
        color=RGBColor(0x1F, 0x38, 0x64))

    # Decorative separator
    _ap(doc, '﷽', sz=18, sb=6, sa=18)

    # Body paragraphs — right-aligned (use LEFT in bidi to get visual right)
    R = WD_ALIGN_PARAGRAPH.LEFT  # visual right in bidi

    p = _ap(doc, '', sz=15, align=R, sb=8, sa=12, ls=1.8)
    r = p.add_run('الحمد لله'); _rf(r, sz=15, bold=True)
    r2 = p.add_run(' رب العالمين، والصلاة والسلام على أشرف المرسلين سيدنا محمد وعلى آله وصحبه أجمعين.')
    _rf(r2, sz=15)

    _ap(doc, 'أتقدم بجزيل الشكر والعرفان إلى الأستاذ(ة) المشرف(ة) على قبوله الإشراف على هذه المذكرة، وعلى توجيهاته القيّمة ونصائحه السديدة التي كان لها الأثر الكبير في إثراء هذا العمل وإتمامه على أحسن وجه.',
        sz=15, align=R, sb=6, sa=12, ls=1.8)

    _ap(doc, 'كما أتوجه بالشكر الجزيل إلى أعضاء لجنة المناقشة المحترمين على تفضلهم بقبول مناقشة هذه المذكرة وتقييمها، وعلى ملاحظاتهم وتوجيهاتهم التي ستُسهم في تحسين هذا العمل.',
        sz=15, align=R, sb=6, sa=12, ls=1.8)

    _ap(doc, 'وأتقدم بالشكر الخالص إلى جميع أساتذة قسم العلوم التجارية الذين رافقونا طوال مشوارنا الجامعي، وأسهموا في تكويننا العلمي والمعرفي.',
        sz=15, align=R, sb=6, sa=12, ls=1.8)

    _ap(doc, 'ولا يفوتني أن أشكر كل من ساعدني من قريب أو بعيد في إنجاز هذا العمل، سواء بالمعلومة أو النصيحة أو الدعم المعنوي.',
        sz=15, align=R, sb=6, sa=12, ls=1.8)

    p = _ap(doc, '', sz=15, bold=True, align=R, sb=6, sa=12, ls=1.8)
    r = p.add_run('جزاكم الله جميعاً خير الجزاء.')
    _rf(r, sz=15, bold=True)

    # Decorative bottom
    _ap(doc, '', sz=12, sb=24, sa=0)
    d2 = doc.add_paragraph(); _deco(d2, '❊', 15, 14)


def build_dedication(doc):
    """الإهداء page — our content from dedication.tex, styled like PDF."""
    pb = doc.add_paragraph()
    pb.paragraph_format.page_break_before = True
    _bidi(pb)

    # Decorative top
    d = doc.add_paragraph(); _deco(d, '❊', 15, 14)

    # Title
    _ap(doc, 'الإهداء', sz=28, bold=True, sb=36, sa=8,
        color=RGBColor(0x1F, 0x38, 0x64))

    # Decorative stars
    _ap(doc, '✻ ✻ ✻ ✻ ✻ ✻ ✻ ✻ ✻ ✻', sz=14, sb=6, sa=18,
        color=RGBColor(0x99, 0x99, 0x99))

    R = WD_ALIGN_PARAGRAPH.LEFT  # visual right in bidi

    _ap(doc, 'إلى من علّمني أن العلم نورٌ يُضيء الدروب...',
        sz=16, align=R, sb=8, sa=14, ls=1.8, italic=True)

    _ap(doc, 'إلى والديّ الكريمين، اللذين كانا سنداً لي في كل خطوة، وغرسا فيّ حب العلم والمثابرة، أسأل الله أن يُطيل في عمرهما ويُديم عليهما الصحة والعافية.',
        sz=15, align=R, sb=6, sa=14, ls=1.8)

    _ap(doc, 'إلى إخوتي وأخواتي الذين شاركوني أفراحي وأحزاني، وكانوا خير عونٍ لي في مسيرتي العلمية.',
        sz=15, align=R, sb=6, sa=14, ls=1.8)

    _ap(doc, 'إلى كل أساتذتي الذين أناروا لي طريق المعرفة، ولم يبخلوا عليّ بعلمهم ونصائحهم القيّمة.',
        sz=15, align=R, sb=6, sa=14, ls=1.8)

    _ap(doc, 'إلى أصدقائي وزملائي الذين رافقوني في هذا المشوار الجامعي، وتقاسمنا معاً لحظات الجد والاجتهاد.',
        sz=15, align=R, sb=6, sa=14, ls=1.8)

    _ap(doc, 'إلى كل من ساهم من قريب أو بعيد في إنجاز هذا العمل المتواضع...',
        sz=15, align=R, sb=6, sa=14, ls=1.8)

    _ap(doc, 'أُهدي هذا العمل.', sz=16, bold=True, align=R, sb=12, sa=12, ls=1.8)

    # Decorative bottom
    _ap(doc, '', sz=12, sb=18, sa=0)
    d2 = doc.add_paragraph(); _deco(d2, '❊', 15, 14)


# ─── body merge ─────────────────────────────────────────────

def copy_body(body_doc, final_doc):
    """Copy body from styled doc starting from المقدمة العامة."""
    # Find start
    body_paras = body_doc.paragraphs
    start_idx = None
    for i, p in enumerate(body_paras):
        if p.style.name == 'Heading 1' and 'المقدمة' in (p.text or ''):
            start_idx = i; break
    if start_idx is None: start_idx = 25
    print(f"  Body starts at paragraph {start_idx}")

    # Find element index
    bb = body_doc.element.body
    kids = list(bb)
    pc = 0; el_start = None
    for idx, ch in enumerate(kids):
        t = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
        if t == 'p':
            if pc == start_idx: el_start = idx; break
            pc += 1
    if el_start is None:
        print("  ERROR: cannot find start element"); return False

    # Map image relationships
    IMAGE_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    HLINK_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'
    rid_map = {}
    for rel in body_doc.part.rels.values():
        try:
            if rel.reltype == IMAGE_REL and not rel.is_external:
                rid_map[rel.rId] = final_doc.part.relate_to(rel.target_part, rel.reltype)
            elif rel.reltype == HLINK_REL and rel.is_external:
                rid_map[rel.rId] = final_doc.part.relate_to(rel.target_ref, rel.reltype, is_external=True)
        except: pass

    r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    fb = final_doc.element.body
    # Find last sectPr
    last_sect = None
    for ch in reversed(list(fb)):
        if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) == 'sectPr':
            last_sect = ch; break

    cnt = 0
    for idx in range(el_start, len(kids)):
        ch = kids[idx]
        t = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
        if t == 'sectPr': continue
        ne = copy.deepcopy(ch)
        for ae in ne.iter():
            for an in [f'{{{r_ns}}}embed', f'{{{r_ns}}}link', f'{{{r_ns}}}id']:
                old = ae.get(an)
                if old and old in rid_map: ae.set(an, rid_map[old])
        if last_sect is not None:
            fb.insert(list(fb).index(last_sect), ne)
        else:
            fb.append(ne)
        cnt += 1
    print(f"  Copied {cnt} elements")
    return True


# ─── section/page numbering ────────────────────────────────

def setup_sections(doc):
    """Insert section break before المقدمة, page numbers start at 1 from body."""
    fb = doc.element.body

    # Find المقدمة by XML
    intro = None
    for ch in fb:
        t = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
        if t != 'p': continue
        pPr = ch.find(qn('w:pPr'))
        if pPr is not None:
            ps = pPr.find(qn('w:pStyle'))
            if ps is not None and ps.get(qn('w:val')) == 'Heading1':
                txt = ''.join((e.text or '') for e in ch.iter(qn('w:t')))
                if 'المقدمة' in txt:
                    intro = ch; break

    if intro is None:
        print("  Warning: Could not find المقدمة"); return

    idx = list(fb).index(intro)
    # Find prev paragraph
    prev = None
    for i in range(idx-1, -1, -1):
        ch = list(fb)[i]
        if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) == 'p':
            prev = ch; break

    if prev:
        pPr = prev.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr'); prev.insert(0, pPr)
        sp = OxmlElement('w:sectPr')
        st = OxmlElement('w:type'); st.set(qn('w:val'), 'nextPage'); sp.append(st)
        ps = OxmlElement('w:pgSz')
        ps.set(qn('w:w'), str(int(21*567))); ps.set(qn('w:h'), str(int(29.7*567)))
        sp.append(ps)
        pm = OxmlElement('w:pgMar')
        pm.set(qn('w:top'), str(int(2.5*567))); pm.set(qn('w:bottom'), str(int(2.5*567)))
        pm.set(qn('w:left'), str(int(2*567))); pm.set(qn('w:right'), str(int(3*567)))
        pm.set(qn('w:header'), '720'); pm.set(qn('w:footer'), '720')
        sp.append(pm)
        sp.append(OxmlElement('w:bidi'))
        pPr.append(sp)

    # Last sectPr = body section
    sects = fb.findall(qn('w:sectPr'))
    if sects:
        ls = sects[-1]
        for o in ls.findall(qn('w:pgSz')): ls.remove(o)
        ps = OxmlElement('w:pgSz')
        ps.set(qn('w:w'), str(int(21*567))); ps.set(qn('w:h'), str(int(29.7*567)))
        ls.append(ps)
        for o in ls.findall(qn('w:pgMar')): ls.remove(o)
        pm = OxmlElement('w:pgMar')
        pm.set(qn('w:top'), str(int(2.5*567))); pm.set(qn('w:bottom'), str(int(2.5*567)))
        pm.set(qn('w:left'), str(int(2*567))); pm.set(qn('w:right'), str(int(3*567)))
        pm.set(qn('w:header'), '720'); pm.set(qn('w:footer'), '720')
        ls.append(pm)
        for o in ls.findall(qn('w:pgNumType')): ls.remove(o)
        pgn = OxmlElement('w:pgNumType'); pgn.set(qn('w:start'), '1'); ls.append(pgn)
        for o in ls.findall(qn('w:bidi')): ls.remove(o)
        ls.append(OxmlElement('w:bidi'))

    # Footer with page number on body section
    sec = doc.sections[-1]
    ftr = sec.footer; ftr.is_linked_to_previous = False
    for p in ftr.paragraphs: p.clear()
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; _bidi(fp)
    r = fp.add_run()
    fb_ = OxmlElement('w:fldChar'); fb_.set(qn('w:fldCharType'), 'begin')
    r._element.append(fb_)
    r2 = fp.add_run()
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve')
    ins.text = ' PAGE '; r2._element.append(ins)
    r3 = fp.add_run()
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    r3._element.append(fe)
    for rr in [r, r2, r3]: rr.font.size = Pt(11); rr.font.name = LF

    # Clear front matter footers
    for i, s in enumerate(doc.sections[:-1]):
        try:
            s.footer.is_linked_to_previous = False
            for p in s.footer.paragraphs: p.clear()
            s.header.is_linked_to_previous = False
            for p in s.header.paragraphs: p.clear()
        except: pass


# ─── main ───────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Building final DOCX (clean, from scratch)")
    print("=" * 60)

    doc = Document()

    # Default style
    n = doc.styles['Normal']
    n.font.name = AF; n.font.size = Pt(14)
    rpr = n.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:cs'), AF); rf.set(qn('w:ascii'), LF); rf.set(qn('w:hAnsi'), LF)
    for o in rpr.findall(qn('w:rFonts')): rpr.remove(o)
    rpr.insert(0, rf)
    pp = n.element.get_or_add_pPr()
    for o in pp.findall(qn('w:bidi')): pp.remove(o)
    pp.append(OxmlElement('w:bidi'))

    # Page setup
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2); sec.right_margin = Cm(3)
    sec.page_height = Cm(29.7); sec.page_width = Cm(21)

    # Remove default empty para
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    print("\n[1/5] Building title page...")
    build_title_page(doc)

    print("[2/5] Building شكر وتقدير...")
    build_acknowledgments(doc)

    print("[3/5] Building الإهداء...")
    build_dedication(doc)

    print("[4/5] Copying body content...")
    body_doc = Document(BODY_DOCX)
    copy_body(body_doc, doc)

    print("[5/5] Setting up page numbering...")
    setup_sections(doc)

    doc.save(OUTPUT)
    kb = os.path.getsize(OUTPUT) / 1024

    # Stats via XML to avoid relationship issues
    h1 = 0
    for ch in doc.element.body:
        t = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
        if t == 'p':
            pPr = ch.find(qn('w:pPr'))
            if pPr is not None:
                ps = pPr.find(qn('w:pStyle'))
                if ps is not None and ps.get(qn('w:val')) == 'Heading1': h1 += 1

    print(f"\n{'='*60}")
    print(f"  SUCCESS! → {OUTPUT}")
    print(f"  Size: {kb:.0f} KB | Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)} | Sections: {len(doc.sections)} | H1: {h1}")
    print(f"{'='*60}")

    shutil.copy2(OUTPUT, os.path.join(BASE, 'main_word.docx'))
    print("  Also copied to main_word.docx")


if __name__ == '__main__':
    main()
