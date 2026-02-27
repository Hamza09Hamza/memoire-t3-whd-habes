#!/usr/bin/env python3
"""
Comprehensive DOCX styler — make main_word.docx look like the LaTeX version.

LaTeX styling:
  - Font: Amiri 14pt (Arabic), Times New Roman (Latin) — we use Traditional Arabic + TNR
  - Margins: top=2.5cm, bottom=2.5cm, right=3cm, left=2cm
  - Line spacing: 1.5
  - Chapter titles: Huge (≈24pt), bold, centered
  - Sections: Large (≈18pt), bold, right-aligned
  - Subsections: large (≈16pt), bold
  - Subsubsections: 14pt, bold
  - Header: chapter name right, page number bottom center
  - Tables: centered, bordered
  - Title page: centered with box around title
  - Dedication/Acknowledgments: right-aligned, decorative spacing
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
INPUT = os.path.join(BASE, 'main_word.docx')
OUTPUT = os.path.join(BASE, 'main_word_styled.docx')

ARABIC_FONT = 'Traditional Arabic'
LATIN_FONT = 'Times New Roman'

# ──────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────

def has_arabic(text):
    return any('\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F' for c in (text or ''))


def set_run_font(run, font_name_ar=ARABIC_FONT, font_name_lat=LATIN_FONT, size_pt=14, bold=None, color=None):
    """Set font properties on a run."""
    rpr = run._element.get_or_add_rPr()

    # Font names
    fonts = rpr.find(qn('w:rFonts'))
    if fonts is None:
        fonts = OxmlElement('w:rFonts')
        rpr.insert(0, fonts)

    text = run.text or ''
    if has_arabic(text):
        fonts.set(qn('w:ascii'), font_name_ar)
        fonts.set(qn('w:hAnsi'), font_name_ar)
    else:
        fonts.set(qn('w:ascii'), font_name_lat)
        fonts.set(qn('w:hAnsi'), font_name_lat)
    fonts.set(qn('w:cs'), font_name_ar)

    # Size
    run.font.size = Pt(size_pt)
    sz_cs = rpr.find(qn('w:szCs'))
    if sz_cs is None:
        sz_cs = OxmlElement('w:szCs')
        rpr.append(sz_cs)
    sz_cs.set(qn('w:val'), str(int(size_pt * 2)))

    # Bold
    if bold is not None:
        run.font.bold = bold
        b_cs = rpr.find(qn('w:bCs'))
        if bold:
            if b_cs is None:
                b_cs = OxmlElement('w:bCs')
                rpr.append(b_cs)
        else:
            if b_cs is not None:
                rpr.remove(b_cs)

    # Color
    if color is not None:
        run.font.color.rgb = color

    # RTL for Arabic runs
    if has_arabic(text):
        rtl = rpr.find(qn('w:rtl'))
        if rtl is None:
            rtl = OxmlElement('w:rtl')
            rpr.append(rtl)


def set_paragraph_format(para, alignment=None, space_before=None, space_after=None,
                         line_spacing=None, line_rule=None, keep_next=False,
                         page_break_before=False, first_line_indent=None):
    """Set paragraph-level formatting."""
    pf = para.paragraph_format

    if alignment is not None:
        pf.alignment = alignment

    if space_before is not None:
        pf.space_before = space_before

    if space_after is not None:
        pf.space_after = space_after

    if line_spacing is not None:
        pf.line_spacing = line_spacing

    if line_rule is not None:
        pf.line_spacing_rule = line_rule

    if keep_next:
        pf.keep_with_next = True

    if page_break_before:
        pf.page_break_before = True

    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent

    # Always set bidi for RTL
    ppr = para._element.get_or_add_pPr()
    bidi = ppr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        ppr.append(bidi)


def set_paragraph_runs_font(para, size_pt=14, bold=None, color=None):
    """Apply font settings to ALL runs in a paragraph."""
    for run in para.runs:
        set_run_font(run, size_pt=size_pt, bold=bold, color=color)


def add_bottom_border(para, sz=4, color='000000'):
    """Add a bottom border to a paragraph (like header rule)."""
    ppr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    ppr.append(pBdr)


def add_top_and_bottom_border(para, sz=12, color='000000'):
    """Add top and bottom borders (for title box effect)."""
    ppr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    ppr.append(pBdr)


def add_box_border(para, sz=12, color='000000'):
    """Add full box border around a paragraph."""
    ppr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    ppr.append(pBdr)


def insert_page_break_before(para):
    """Force page break before this paragraph."""
    set_paragraph_format(para, page_break_before=True)


def set_cell_shading(cell, color='D9E2F3'):
    """Set background color on a table cell."""
    tc = cell._element
    tc_pr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tc_pr.append(shading)


def set_cell_borders(cell, sz=4, color='000000'):
    """Set borders on a single table cell."""
    tc = cell._element
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_table_style(table):
    """Style a table: centered, borders, header row shading."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style all cells
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell, sz=4, color='000000')

            # Vertical alignment
            tc_pr = cell._element.get_or_add_tcPr()
            v_align = OxmlElement('w:vAlign')
            v_align.set(qn('w:val'), 'center')
            tc_pr.append(v_align)

            for para in cell.paragraphs:
                set_paragraph_format(para,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     space_before=Pt(2),
                                     space_after=Pt(2),
                                     line_spacing=1.15)
                set_paragraph_runs_font(para, size_pt=11, bold=None)

            # Header row: bold + shading
            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for para in cell.paragraphs:
                    set_paragraph_runs_font(para, size_pt=11, bold=True)


def add_header_footer(section, header_text=''):
    """Add header with text on right and page number in footer center."""
    # Header
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()
    hp.text = ''
    # We can't dynamically set chapter name, but we add a placeholder format
    # Set header paragraph formatting
    set_paragraph_format(hp, alignment=WD_ALIGN_PARAGRAPH.LEFT)  # bidi: LEFT→visually RIGHT
    # Add bottom border (like \headrulewidth)
    add_bottom_border(hp, sz=4, color='808080')

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    fp.text = ''
    set_paragraph_format(fp, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Add PAGE field
    run = fp.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_char_begin)

    run2 = fp.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._element.append(instr)

    run3 = fp.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run3._element.append(fld_char_end)

    for r in [run, run2, run3]:
        set_run_font(r, size_pt=11)


# ──────────────────────────────────────────────────────────
# Main styling
# ──────────────────────────────────────────────────────────

def main():
    print("Loading main_word.docx...")
    doc = Document(INPUT)

    # ── 1. DOCUMENT-LEVEL SETTINGS ──────────────────────────
    print("[1/7] Setting page margins, line spacing, section properties...")

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(3)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)

        # Set document bidi
        sect_pr = section._sectPr
        bidi = sect_pr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            sect_pr.append(bidi)

    # ── 2. STYLE DEFINITIONS ──────────────────────────────────
    print("[2/7] Configuring document styles (Normal, Headings)...")

    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = ARABIC_FONT
    normal.font.size = Pt(14)
    nrpr = normal.element.get_or_add_rPr()
    nfonts = OxmlElement('w:rFonts')
    nfonts.set(qn('w:cs'), ARABIC_FONT)
    nfonts.set(qn('w:ascii'), LATIN_FONT)
    nfonts.set(qn('w:hAnsi'), LATIN_FONT)
    # Remove old rFonts if any
    for old in nrpr.findall(qn('w:rFonts')):
        nrpr.remove(old)
    nrpr.insert(0, nfonts)

    nsz = OxmlElement('w:szCs')
    nsz.set(qn('w:val'), '28')
    for old in nrpr.findall(qn('w:szCs')):
        nrpr.remove(old)
    nrpr.append(nsz)

    nppr = normal.element.get_or_add_pPr()
    # Line spacing 1.5
    spacing = nppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        nppr.append(spacing)
    spacing.set(qn('w:line'), '360')  # 360 twips = 1.5 lines (240 twips = single)
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '120')   # 6pt after
    spacing.set(qn('w:before'), '0')

    # RTL/bidi on Normal
    for old in nppr.findall(qn('w:bidi')):
        nppr.remove(old)
    bidi_n = OxmlElement('w:bidi')
    nppr.append(bidi_n)

    # Justify
    for old in nppr.findall(qn('w:jc')):
        nppr.remove(old)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    nppr.append(jc)

    # Heading styles configuration
    # NOTE: With w:bidi set, Word swaps left/right semantics.
    # So 'left' in XML = visually RIGHT for RTL paragraphs.
    heading_configs = {
        'Heading 1': {'size': 24, 'color': '1F3864', 'space_before': 0, 'space_after': 480, 'align': 'center', 'page_break': True},
        'Heading 2': {'size': 18, 'color': '2E4057', 'space_before': 360, 'space_after': 200, 'align': 'left', 'page_break': False},
        'Heading 3': {'size': 16, 'color': '374151', 'space_before': 240, 'space_after': 120, 'align': 'left', 'page_break': False},
        'Heading 4': {'size': 14, 'color': '4B5563', 'space_before': 200, 'space_after': 100, 'align': 'left', 'page_break': False},
    }

    for style_name, cfg in heading_configs.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = ARABIC_FONT
        style.font.size = Pt(cfg['size'])
        style.font.bold = True

        if cfg['color']:
            style.font.color.rgb = RGBColor.from_string(cfg['color'])

        # rPr
        rpr = style.element.get_or_add_rPr()
        for old in rpr.findall(qn('w:rFonts')):
            rpr.remove(old)
        hfonts = OxmlElement('w:rFonts')
        hfonts.set(qn('w:cs'), ARABIC_FONT)
        hfonts.set(qn('w:ascii'), LATIN_FONT)
        hfonts.set(qn('w:hAnsi'), LATIN_FONT)
        rpr.insert(0, hfonts)

        for old in rpr.findall(qn('w:szCs')):
            rpr.remove(old)
        hsz = OxmlElement('w:szCs')
        hsz.set(qn('w:val'), str(cfg['size'] * 2))
        rpr.append(hsz)

        for old in rpr.findall(qn('w:rtl')):
            rpr.remove(old)
        hrtl = OxmlElement('w:rtl')
        rpr.append(hrtl)

        for old in rpr.findall(qn('w:bCs')):
            rpr.remove(old)
        hbcs = OxmlElement('w:bCs')
        rpr.append(hbcs)

        # pPr
        ppr = style.element.get_or_add_pPr()

        for old in ppr.findall(qn('w:bidi')):
            ppr.remove(old)
        hbidi = OxmlElement('w:bidi')
        ppr.append(hbidi)

        for old in ppr.findall(qn('w:jc')):
            ppr.remove(old)
        hjc = OxmlElement('w:jc')
        hjc.set(qn('w:val'), cfg['align'])
        ppr.append(hjc)

        # Spacing
        for old in ppr.findall(qn('w:spacing')):
            ppr.remove(old)
        hspacing = OxmlElement('w:spacing')
        hspacing.set(qn('w:before'), str(cfg['space_before']))
        hspacing.set(qn('w:after'), str(cfg['space_after']))
        hspacing.set(qn('w:line'), '360')
        hspacing.set(qn('w:lineRule'), 'auto')
        ppr.append(hspacing)

        # Keep with next
        kwn = ppr.find(qn('w:keepNext'))
        if kwn is None:
            kwn = OxmlElement('w:keepNext')
            ppr.append(kwn)

    # ── 3. TITLE PAGE ─────────────────────────────────────────
    print("[3/7] Styling title page...")

    paras = doc.paragraphs
    total = len(paras)

    # Title page paragraphs: 0-4
    # 0: Institution name
    if total > 0:
        p = paras[0]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(36), space_after=Pt(48))
        set_paragraph_runs_font(p, size_pt=16, bold=True)

    # 1: Thesis type + degree
    if total > 1:
        p = paras[1]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(48), space_after=Pt(24))
        set_paragraph_runs_font(p, size_pt=16, bold=True)

    # 2: Host institution + students (multi-line paragraph)
    if total > 2:
        p = paras[2]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(24), space_after=Pt(12))
        set_paragraph_runs_font(p, size_pt=14, bold=True)

    # 3: Thesis title - this is the KEY visual element, should have box border
    # Actually looking at the structure: para 0 = institution, para 1 = merged text with title type,
    # Let me find the one with the actual thesis title
    # From the check: para 0 = المعهد, para 1 = مذكرة + شهادة, para 2 = مؤسسة + students
    # para 3 = تأطير, para 4 = دفعة
    # The thesis title "تحديات وكالات الأسفار في ظل المنافسة" seems embedded
    # Let me check if it's in para 1 or separate

    # Actually from the pandoc output, the title's fcolorbox content got merged.
    # Let me find it by searching
    thesis_title_idx = None
    for i in range(min(10, total)):
        t = paras[i].text.strip()
        if 'تحديات وكالات الأسفار' in t and len(t) < 200:
            thesis_title_idx = i
            break

    if thesis_title_idx is not None:
        # It's likely embedded in para 1 which has multi-line content
        # We need to handle differently based on structure
        pass

    # Style paras 3 and 4
    if total > 3:
        p = paras[3]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(12), space_after=Pt(12))
        set_paragraph_runs_font(p, size_pt=14, bold=True)

    if total > 4:
        p = paras[4]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(48), space_after=Pt(0))
        set_paragraph_runs_font(p, size_pt=16, bold=True)

    # Now let's try to find and style the title text within para 1
    # Para 1 contains the thesis type + title in merged text
    if total > 1:
        p1_text = paras[1].text
        if 'تحديات وكالات الأسفار' in p1_text or 'مذكــ' in p1_text:
            # This is the big title paragraph, add box border effect
            add_top_and_bottom_border(paras[1], sz=8, color='000000')

    # ── 4. FRONT MATTER (dedication, acknowledgments, abstract) ──
    print("[4/7] Styling front matter (dedication, acknowledgments, abstract)...")

    # Dedication: paragraphs 5-12
    # Para 5 = "الإهداء" title
    if total > 5:
        p = paras[5]
        insert_page_break_before(p)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(72), space_after=Pt(48))
        set_paragraph_runs_font(p, size_pt=26, bold=True, color=RGBColor(31, 56, 100))

    # Dedication body: 6-12 (bidi swaps LEFT→visually RIGHT)
    for i in range(6, min(13, total)):
        p = paras[i]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(6), space_after=Pt(12),
                             line_spacing=1.8, first_line_indent=Cm(0))
        set_paragraph_runs_font(p, size_pt=15)

    # Acknowledgments: paragraphs 13-19
    # Para 13 = "شكر وتقدير" title
    if total > 13:
        p = paras[13]
        insert_page_break_before(p)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(72), space_after=Pt(36))
        set_paragraph_runs_font(p, size_pt=26, bold=True, color=RGBColor(31, 56, 100))

    # Acknowledgments body: 14-19 (bidi swaps LEFT→visually RIGHT)
    for i in range(14, min(20, total)):
        p = paras[i]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(6), space_after=Pt(12),
                             line_spacing=1.8)
        set_paragraph_runs_font(p, size_pt=15)

    # Abstract: paragraphs 20-24
    # Para 20 = "ملخص الدراسة" title
    if total > 20:
        p = paras[20]
        insert_page_break_before(p)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(72), space_after=Pt(36))
        set_paragraph_runs_font(p, size_pt=26, bold=True, color=RGBColor(31, 56, 100))

    # Abstract body: 21-24
    for i in range(21, min(25, total)):
        p = paras[i]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=Pt(4), space_after=Pt(8),
                             line_spacing=1.8)
        set_paragraph_runs_font(p, size_pt=14)

    # ── 5. BODY CONTENT ─────────────────────────────────────
    print("[5/7] Styling body content (headings, paragraphs, lists)...")

    # Heading 1 indices: 25, 82, 358, 600, 925
    heading1_indices = []
    for i, p in enumerate(paras):
        if p.style.name == 'Heading 1':
            heading1_indices.append(i)

    # Style all Heading 1 (chapters)
    for idx in heading1_indices:
        p = paras[idx]
        insert_page_break_before(p)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(0), space_after=Pt(36))
        set_paragraph_runs_font(p, size_pt=24, bold=True, color=RGBColor(31, 56, 100))
        add_bottom_border(p, sz=6, color='1F3864')

    # Style all Heading 2 (sections) — bidi swaps LEFT→visually RIGHT
    for i, p in enumerate(paras):
        if p.style.name == 'Heading 2':
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=Pt(24), space_after=Pt(12),
                                 keep_next=True)
            set_paragraph_runs_font(p, size_pt=18, bold=True, color=RGBColor(46, 64, 87))

    # Style all Heading 3 (subsections) — bidi swaps LEFT→visually RIGHT
    for i, p in enumerate(paras):
        if p.style.name == 'Heading 3':
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=Pt(18), space_after=Pt(8),
                                 keep_next=True)
            set_paragraph_runs_font(p, size_pt=16, bold=True, color=RGBColor(55, 65, 81))

    # Style all Heading 4 (subsubsections) — bidi swaps LEFT→visually RIGHT
    for i, p in enumerate(paras):
        if p.style.name == 'Heading 4':
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=Pt(14), space_after=Pt(6),
                                 keep_next=True)
            set_paragraph_runs_font(p, size_pt=14, bold=True, color=RGBColor(75, 85, 99))

    # Style all Normal body paragraphs (after front matter, idx >= 25)
    for i in range(25, total):
        p = paras[i]
        if p.style.name == 'Normal':
            text = p.text.strip()
            if not text:
                continue

            # Detect list items (bullets/numbers) — they often start with specific patterns
            is_list = False
            for marker in ['--', '►', '•']:
                if text.startswith(marker):
                    is_list = True
                    break

            if is_list:
                set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     space_before=Pt(2), space_after=Pt(2),
                                     line_spacing=1.5,
                                     first_line_indent=Cm(0))
                set_paragraph_runs_font(p, size_pt=14)
            else:
                # Regular paragraph — justified with 1.5 spacing
                set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                     space_before=Pt(0), space_after=Pt(6),
                                     line_spacing=1.5)
                set_paragraph_runs_font(p, size_pt=14)

    # ── 6. TABLES ────────────────────────────────────────────
    print("[6/7] Styling tables...")

    for table in doc.tables:
        set_table_style(table)

    # ── 7. HEADERS & FOOTERS ─────────────────────────────────
    print("[7/7] Adding headers and footers...")

    for section in doc.sections:
        add_header_footer(section)

    # ── SAVE ──────────────────────────────────────────────────
    doc.save(OUTPUT)
    size_kb = os.path.getsize(OUTPUT) / 1024
    print(f'\n{"=" * 60}')
    print(f'  SUCCESS! Styled document saved as:')
    print(f'  {OUTPUT}')
    print(f'  Size: {size_kb:.0f} KB')
    print(f'{"=" * 60}')


if __name__ == '__main__':
    main()
