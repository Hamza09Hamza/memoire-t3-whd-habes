"""Post-process the DOCX to ensure Arabic font and RTL are applied everywhere."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('main_word.docx')

def set_paragraph_rtl_and_font(paragraph, font_name='Traditional Arabic', font_size_pt=14):
    """Ensure a paragraph has RTL direction and Arabic font."""
    # Set paragraph-level bidi
    ppr = paragraph._element.get_or_add_pPr()
    bidi = ppr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        ppr.append(bidi)

    # Set each run's font
    for run in paragraph.runs:
        rpr = run._element.get_or_add_rPr()
        
        # Set complex script font
        existing_fonts = rpr.find(qn('w:rFonts'))
        if existing_fonts is None:
            existing_fonts = OxmlElement('w:rFonts')
            rpr.append(existing_fonts)
        existing_fonts.set(qn('w:cs'), font_name)
        
        # If the text is primarily Latin/English, use Times New Roman
        text = run.text or ''
        has_arabic = any('\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F' for c in text)
        if has_arabic:
            existing_fonts.set(qn('w:ascii'), font_name)
            existing_fonts.set(qn('w:hAnsi'), font_name)
        else:
            existing_fonts.set(qn('w:ascii'), 'Times New Roman')
            existing_fonts.set(qn('w:hAnsi'), 'Times New Roman')
        
        # Set RTL for runs with Arabic text
        if has_arabic:
            rtl = rpr.find(qn('w:rtl'))
            if rtl is None:
                rtl = OxmlElement('w:rtl')
                rpr.append(rtl)

        # Set cs font size
        sz_cs = rpr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = OxmlElement('w:szCs')
            rpr.append(sz_cs)
        sz_cs.set(qn('w:val'), str(font_size_pt * 2))

# Process all paragraphs
for para in doc.paragraphs:
    style = para.style.name
    if 'Heading 1' in style:
        set_paragraph_rtl_and_font(para, font_size_pt=22)
    elif 'Heading 2' in style:
        set_paragraph_rtl_and_font(para, font_size_pt=18)
    elif 'Heading 3' in style:
        set_paragraph_rtl_and_font(para, font_size_pt=16)
    elif 'Heading 4' in style:
        set_paragraph_rtl_and_font(para, font_size_pt=14)
    else:
        set_paragraph_rtl_and_font(para, font_size_pt=14)

# Process table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                set_paragraph_rtl_and_font(para, font_size_pt=12)

# Set document-level RTL in section properties
for section in doc.sections:
    sectPr = section._sectPr
    bidi = sectPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)

doc.save('main_word.docx')
print("Post-processing complete! Arabic fonts and RTL applied to all content.")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
