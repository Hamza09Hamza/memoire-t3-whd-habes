from docx import Document
from docx.oxml.ns import qn

doc = Document('main_word_styled.docx')

# Check margins
s = doc.sections[0]
print(f"Margins: top={s.top_margin.cm:.1f}cm, bottom={s.bottom_margin.cm:.1f}cm, "
      f"left={s.left_margin.cm:.1f}cm, right={s.right_margin.cm:.1f}cm")

# Check Normal style
normal = doc.styles['Normal']
print(f"\nNormal style: font={normal.font.name}, size={normal.font.size}")

# Check first few paragraphs alignment & font
print("\n--- Title page paragraphs ---")
for i in range(5):
    p = doc.paragraphs[i]
    text = p.text.strip()[:50]
    align = str(p.paragraph_format.alignment)
    font_size = None
    if p.runs:
        font_size = p.runs[0].font.size
    print(f"  [{i}] align={align}, font_size={font_size}, text={text}")

print("\n--- Dedication title ---")
p = doc.paragraphs[5]
print(f"  text={p.text.strip()[:30]}, align={p.paragraph_format.alignment}, "
      f"page_break={p.paragraph_format.page_break_before}")
if p.runs:
    r = p.runs[0]
    print(f"  font_size={r.font.size}, bold={r.font.bold}, color={r.font.color.rgb}")

print("\n--- Heading 1 samples ---")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1':
        text = p.text.strip()[:40]
        align = str(p.paragraph_format.alignment)
        pb = p.paragraph_format.page_break_before
        fs = None
        if p.runs:
            fs = p.runs[0].font.size
        print(f"  [{i}] align={align}, page_break={pb}, font={fs}, text={text}")

print("\n--- Tables ---")
for i, table in enumerate(doc.tables[:3]):
    print(f"  Table {i+1}: alignment={table.alignment}")
    # Check first row shading
    cell = table.cell(0, 0)
    tc_pr = cell._element.find(qn('w:tcPr'))
    shd = tc_pr.find(qn('w:shd')) if tc_pr is not None else None
    fill = shd.get(qn('w:fill')) if shd is not None else None
    print(f"    Header shading: {fill}")
    print(f"    First cell text: {cell.text.strip()[:30]}")

# Check header/footer
print("\n--- Header/Footer ---")
for si, sec in enumerate(doc.sections):
    h = sec.header
    f = sec.footer
    h_text = h.paragraphs[0].text if h.paragraphs else '(none)'
    f_has_field = any('PAGE' in (r.text or '') for fp in f.paragraphs for r in fp.runs) if f.paragraphs else False
    # Check for field code in XML
    f_xml = f._element.xml if f else ''
    has_page = 'PAGE' in f_xml
    print(f"  Section {si}: header='{h_text[:30]}', footer_has_page_num={has_page}")
