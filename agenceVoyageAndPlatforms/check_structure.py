from docx import Document
doc = Document('main_word.docx')

# Print the first 60 paragraphs with index, style, and text
for i, p in enumerate(doc.paragraphs[:80]):
    text = p.text.strip()
    if text:
        preview = text[:70].replace('\n', ' | ')
        print(f'{i:4d} [{p.style.name:12s}] {preview}')
    else:
        print(f'{i:4d} [{p.style.name:12s}] (empty)')
    if i == 79:
        break

print('\n\n--- Heading 1 positions ---')
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1':
        print(f'{i:4d} {p.text.strip()[:60]}')

print('\n--- Total ---')
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
