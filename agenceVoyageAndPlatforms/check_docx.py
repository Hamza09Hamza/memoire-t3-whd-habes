from docx import Document
doc = Document('main_word.docx')
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
print()

# Show first 15 paragraphs with their styles
for i, p in enumerate(doc.paragraphs[:20]):
    text = p.text.strip()
    if text:
        style = p.style.name
        preview = text[:80] + ('...' if len(text) > 80 else '')
        print(f'[{style}] {preview}')

print()
print('--- Middle content sample ---')
mid = len(doc.paragraphs) // 2
for p in doc.paragraphs[mid:mid+5]:
    text = p.text.strip()
    if text:
        preview = text[:80] + ('...' if len(text) > 80 else '')
        print(f'[{p.style.name}] {preview}')

print()
print('--- Tables check ---')
for i, table in enumerate(doc.tables[:3]):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.cell(0,0).text.strip()[:50]
    print(f'Table {i+1}: {rows} rows x {cols} cols | First cell: {first_cell}')
